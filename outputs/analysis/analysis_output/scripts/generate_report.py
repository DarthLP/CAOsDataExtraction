#!/usr/bin/env python3
"""
generate_report.py

Description:
    Generates a Word report that compiles key plots and metrics from the Part 1 and Part 2 analyses.
    The report is structured by content themes (coverage, gaps & renewals, seasonality, salary tables,
    benefits, missingness, and activity over time), not by script parts. It embeds plots and summarizes
    key figures from the produced Excel tables.

Usage:
    python generate_report.py \
        --outdir "/absolute/path/to/analysis_output" \
        --output "/absolute/path/to/analysis_output/CAO_analysis_report.docx"

Notes:
    - Expects plots under: {outdir}/plots/part1 and {outdir}/plots/part2
    - Expects tables under: {outdir}/tables/part1, {outdir}/tables/part1/details, {outdir}/tables/part2
    - Uses python-docx to generate a .docx file
    - Console output is minimal: prints start and completion messages only
"""

from __future__ import annotations

import argparse
import os
from datetime import datetime
from typing import Dict, Optional

import pandas as pd

from docx import Document
from docx.shared import Inches


def path(*parts: str) -> str:
    return os.path.join(*parts)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_image(doc: Document, image_path: str, caption: Optional[str] = None, width_in: float = 6.0) -> None:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(width_in))
        if caption:
            p = doc.add_paragraph()
            p.add_run(caption).italic = True


def try_read_excel(excel_path: str, sheet_name: int | str = 0) -> Optional[pd.DataFrame]:
    try:
        return pd.read_excel(excel_path, sheet_name=sheet_name)
    except Exception:
        return None


def compute_part1_summaries(tables_dir_part1: str) -> Dict[str, object]:
    results: Dict[str, object] = {}

    coverage_path = path(tables_dir_part1, "cao_coverage_summary.xlsx")
    coverage_df = try_read_excel(coverage_path, sheet_name="Data")
    if coverage_df is None:
        coverage_df = try_read_excel(coverage_path)
    if coverage_df is not None and not coverage_df.empty:
        results["num_caos"] = int(coverage_df["cao_number"].nunique()) if "cao_number" in coverage_df.columns else len(coverage_df)
        results["fully_covered_caos"] = int(coverage_df.get("is_fully_covered", pd.Series(dtype=bool)).sum()) if "is_fully_covered" in coverage_df.columns else None
        results["avg_coverage_months"] = float(coverage_df.get("coverage_months", pd.Series(dtype=float)).mean()) if "coverage_months" in coverage_df.columns else None
        results["avg_files_per_cao"] = float(coverage_df.get("num_files", pd.Series(dtype=float)).mean()) if "num_files" in coverage_df.columns else None

    renew_overall_path = path(tables_dir_part1, "renewal_gap_stats_overall.xlsx")
    renew_overall_df = try_read_excel(renew_overall_path, sheet_name="Data")
    if renew_overall_df is None:
        renew_overall_df = try_read_excel(renew_overall_path)
    if renew_overall_df is not None and not renew_overall_df.empty:
        # Expect rows for 'all' and 'positive_only'
        results["renewal_overall"] = renew_overall_df

    renew_per_cao_path = path(tables_dir_part1, "renewal_gap_stats_per_cao.xlsx")
    renew_per_cao_df = try_read_excel(renew_per_cao_path, sheet_name="Data")
    if renew_per_cao_df is None:
        renew_per_cao_df = try_read_excel(renew_per_cao_path)
    if renew_per_cao_df is not None and not renew_per_cao_df.empty:
        results["renewal_per_cao"] = renew_per_cao_df

    return results


def compute_part2_summaries(tables_dir_part2: str) -> Dict[str, object]:
    results: Dict[str, object] = {}

    salary_presence_path = path(tables_dir_part2, "cao_salary_table_presence.xlsx")
    salary_presence_df = try_read_excel(salary_presence_path, sheet_name="Data")
    if salary_presence_df is None:
        salary_presence_df = try_read_excel(salary_presence_path)
    if salary_presence_df is not None and not salary_presence_df.empty:
        results["salary_presence"] = salary_presence_df

    completeness_year_path = path(tables_dir_part2, "salary_table_completeness_by_year.xlsx")
    completeness_year_df = try_read_excel(completeness_year_path, sheet_name="Data")
    if completeness_year_df is None:
        completeness_year_df = try_read_excel(completeness_year_path)
    if completeness_year_df is not None and not completeness_year_df.empty:
        results["completeness_by_year"] = completeness_year_df

    benefits_presence_path = path(tables_dir_part2, "benefit_presence_summary.xlsx")
    benefits_presence_df = try_read_excel(benefits_presence_path, sheet_name="Data")
    if benefits_presence_df is None:
        benefits_presence_df = try_read_excel(benefits_presence_path)
    if benefits_presence_df is not None and not benefits_presence_df.empty:
        results["benefits_presence"] = benefits_presence_df

    missingness_path = path(tables_dir_part2, "missingness_by_cao.xlsx")
    missingness_df = try_read_excel(missingness_path, sheet_name="Data")
    if missingness_df is None:
        missingness_df = try_read_excel(missingness_path)
    if missingness_df is not None and not missingness_df.empty:
        results["missingness"] = missingness_df

    richness_by_sector_path = path(tables_dir_part2, "benefit_richness_by_sector.xlsx")
    richness_by_sector_df = try_read_excel(richness_by_sector_path, sheet_name="Data")
    if richness_by_sector_df is None:
        richness_by_sector_df = try_read_excel(richness_by_sector_path)
    if richness_by_sector_df is not None and not richness_by_sector_df.empty:
        results["richness_by_sector"] = richness_by_sector_df

    return results


def add_coverage_section(doc: Document, plots_dir_part1: str, p1: Dict[str, object]) -> None:
    add_heading(doc, "Coverage and Periods", level=1)
    add_paragraph(doc, "Coverage is defined per CAO as the union of all validity periods we found in the source files. The earliest start date and latest expiry date establish the outer window. Inside that window there may be internal gaps (days where no period was in force) or overlaps (two or more periods simultaneously covering the same time). This section summarizes those high‑level time bounds and how many files contribute to each CAO.")
    if "num_caos" in p1:
        add_paragraph(doc, f"Total CAOs analyzed: {p1['num_caos']}")
    if "fully_covered_caos" in p1 and p1["fully_covered_caos"] is not None:
        add_paragraph(doc, f"CAOs fully covered without internal gaps: {p1['fully_covered_caos']}")
    if "avg_coverage_months" in p1 and p1["avg_coverage_months"] is not None:
        add_paragraph(doc, f"Average coverage window (earliest start to latest end): {p1['avg_coverage_months']:.1f} months")
    if "avg_files_per_cao" in p1 and p1["avg_files_per_cao"] is not None:
        add_paragraph(doc, f"Average number of files per CAO: {p1['avg_files_per_cao']:.2f}")
    add_paragraph(doc, "Interpretation guidance: a long coverage window with few files suggests sparse renewals or long multi‑year agreements; a short window with many files suggests frequent renewals, addenda, or multiple document variants.")

    add_image(doc, path(plots_dir_part1, "hist_earliest_start_years.png"), "Earliest start years distribution")
    add_image(doc, path(plots_dir_part1, "hist_latest_expiry_years.png"), "Latest expiry years distribution")
    add_image(doc, path(plots_dir_part1, "hist_num_files_per_cao.png"), "Files per CAO distribution")
    add_image(doc, path(plots_dir_part1, "scatter_files_vs_coverage_months.png"), "Files vs coverage length (months)")


def add_gaps_renewals_section(doc: Document, plots_dir_part1: str, p1: Dict[str, object]) -> None:
    add_heading(doc, "Gaps and Renewals", level=1)
    add_paragraph(doc, "Renewal gap = next period's start date minus previous period's end date after sorting by start date within a CAO. We distinguish: (a) negative gaps = overlaps (the next starts before the previous ends), (b) zero gaps = contiguous renewals (back‑to‑back), and (c) positive gaps = uncovered time. Only positive gaps represent true holes in coverage.")
    add_paragraph(doc, "Large negative gaps typically arise when published periods use placeholder end dates (e.g., 2028‑12‑31) that were superseded early. Use the positive‑only statistics to understand actual breaks; use the 'all transitions' statistics to understand administrative cadence including overlaps.")

    # Overall renewal stats
    renew_overall_df = p1.get("renewal_overall")
    if isinstance(renew_overall_df, pd.DataFrame) and not renew_overall_df.empty:
        try:
            row_all = renew_overall_df.loc[renew_overall_df["subset"] == "all"].iloc[0]
            row_pos = renew_overall_df.loc[renew_overall_df["subset"] == "positive_only"].iloc[0]
            add_paragraph(doc, f"Renewal gaps (all transitions): mean {row_all['overall_avg_gap_months']:.2f} months, median {row_all['overall_median_gap_months']:.2f} months.")
            add_paragraph(doc, f"Renewal gaps (positive-only): mean {row_pos['overall_avg_gap_months']:.2f} months, median {row_pos['overall_median_gap_months']:.2f} months.")
        except Exception:
            pass

    add_image(doc, path(plots_dir_part1, "hist_renewal_gap_lengths_overall.png"), "Renewal gap lengths histogram (months)")


def add_seasonality_section(doc: Document, plots_dir_part1: str) -> None:
    add_heading(doc, "Seasonality of Dates", level=1)
    add_paragraph(doc, "We show the month of all period starts and expiries across the corpus. This can reveal preferred renewal months (e.g., January or July) or administrative clustering (e.g., expiries at year‑end). These charts are descriptive; they do not account for period length or document weight.")
    add_image(doc, path(plots_dir_part1, "seasonality_all_start_months.png"), "All start months")
    add_image(doc, path(plots_dir_part1, "seasonality_all_expiry_months.png"), "All expiry months")


def add_salary_section(doc: Document, plots_dir_part2: str, p2: Dict[str, object]) -> None:
    add_heading(doc, "Salary Tables", level=1)
    add_paragraph(doc, "We count salary tables per file by checking filled columns salary_1…salary_7. If all seven are filled and more_salaries is 'yes', we mark '8+' (internally 8). Completeness over time is the share of files per start year that contain ≥1 salary table. This reflects the presence of structured pay information, not its depth or correctness.")
    add_paragraph(doc, "Reading tips: spikes in the distribution at 1–2 tables often reflect single‑table salary structures or one update table alongside an existing scale. '8+' indicates rich multi‑table structures. Over‑time completeness can shift as extraction quality or document formats evolve.")
    add_image(doc, path(plots_dir_part2, "hist_salary_tables_per_file_percent_and_count.png"), "Salary tables per file: % and count (aligned axes)")
    add_image(doc, path(plots_dir_part2, "line_salary_completeness_over_time.png"), "% of files with ≥1 salary table by start year, with file counts")

    comp_df = p2.get("completeness_by_year")
    if isinstance(comp_df, pd.DataFrame) and not comp_df.empty:
        try:
            last = comp_df.sort_values("start_year").dropna(subset=["pct_with_salary"]).iloc[-1]
            add_paragraph(doc, f"Latest year completeness: {last['start_year']}: {last['pct_with_salary']:.1f}% of files with ≥1 salary table (n={int(last['num_files'])}).")
        except Exception:
            pass


def add_benefits_section(doc: Document, plots_dir_part2: str, p2: Dict[str, object]) -> None:
    add_heading(doc, "Benefits", level=1)
    add_paragraph(doc, "Benefit presence is detected when any mapped column for a topic is non‑empty in a file. Mapping includes: pension (pension, retire), leave (vacation, maternity, vakantie, verlof), termination (term_*, ontslag, beëindiging, probation), overtime (overtime, shift compensation, max/min hours), training, and homeoffice. Presence indicates the topic is mentioned with data, not necessarily that it is exhaustive or standardized.")
    add_paragraph(doc, "Use the prevalence‑over‑time chart to see adoption trends (e.g., rise of homeoffice). Presence can be conservative for sparse entries and liberal for verbose narrative; interpret comparatively across topics and years rather than as absolute compliance.")
    add_image(doc, path(plots_dir_part2, "hist_benefits_percent_and_count.png"), "Benefits presence: % (left) and counts (right scale)")
    add_image(doc, path(plots_dir_part2, "line_benefit_prevalence_over_time.png"), "Benefit prevalence over time (% of files)")

    ben_df = p2.get("benefits_presence")
    if isinstance(ben_df, pd.DataFrame) and not ben_df.empty:
        try:
            top = ben_df.sort_values("pct", ascending=False).head(3)
            items = ", ".join([f"{r['benefit']}: {r['pct']:.1f}%" for _, r in top.iterrows()])
            add_paragraph(doc, f"Top benefits by presence: {items}.")
        except Exception:
            pass


def add_missingness_section(doc: Document, plots_dir_part2: str, p2: Dict[str, object]) -> None:
    add_heading(doc, "Missingness by CAO", level=1)
    add_paragraph(doc, "Missingness is computed at file level per CAO: salary is missing when the file has zero salary tables; a benefit is missing when none of its mapped columns are filled. We then average within each CAO. High missingness can reflect either true absence in the documents or extraction gaps; use alongside presence and completeness metrics.")
    add_image(doc, path(plots_dir_part2, "heatmap_missingness_by_cao.png"), "Missingness heatmap (% missing by CAO)")


def add_activity_over_time_section(doc: Document, plots_dir_part2: str) -> None:
    add_heading(doc, "Files Over Time", level=1)
    add_paragraph(doc, "We compare counts of files by earliest start year and by end year. Divergence between the lines can indicate long periods (many end in later years) or back‑dated agreements (starts cluster earlier). This provides context for interpreting completeness/prevalence trends.")
    add_image(doc, path(plots_dir_part2, "line_num_files_per_year_start_vs_end.png"), "Files per year: start vs end")


def build_report(outdir: str, output_docx: str) -> None:
    plots_dir_part1 = path(outdir, "plots", "part1")
    plots_dir_part2 = path(outdir, "plots", "part2")
    tables_dir_part1 = path(outdir, "tables", "part1")
    tables_dir_part1_details = path(tables_dir_part1, "details")
    tables_dir_part2 = path(outdir, "tables", "part2")

    # Compute summaries
    p1 = compute_part1_summaries(tables_dir_part1)
    p2 = compute_part2_summaries(tables_dir_part2)

    # Build document
    doc = Document()
    doc.add_heading("CAO Analysis Report", 0)
    add_paragraph(doc, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    add_paragraph(doc, "This report consolidates date coverage, renewal dynamics, salary information, and benefits presence across CAOs. It explains how each metric is constructed and how to read the visualizations. We avoid repeating labels already visible in charts; instead we focus on definitions, caveats, and how to connect the pieces into a coherent view.")

    add_coverage_section(doc, plots_dir_part1, p1)
    add_gaps_renewals_section(doc, plots_dir_part1, p1)
    add_seasonality_section(doc, plots_dir_part1)
    add_salary_section(doc, plots_dir_part2, p2)
    add_benefits_section(doc, plots_dir_part2, p2)
    add_missingness_section(doc, plots_dir_part2, p2)
    add_activity_over_time_section(doc, plots_dir_part2)

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    doc.save(output_docx)


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate a Word report from analysis outputs")
    parser.add_argument("--outdir", required=True, help="Absolute path to output root directory (analysis_output)")
    parser.add_argument("--output", required=False, default=None, help="Absolute path to the .docx to write")
    args = parser.parse_args()

    output_docx = args.output or path(args.outdir, "CAO_analysis_report.docx")

    print("Building Word report...")
    build_report(args.outdir, output_docx)
    print(f"Saved report to: {output_docx}")


if __name__ == "__main__":
    main()


